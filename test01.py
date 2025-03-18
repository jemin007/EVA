import os
from typing import List, Dict, Any, Set
from fastapi import FastAPI, HTTPException, Depends, Request,Security
from fastapi.responses import FileResponse
from pydantic import BaseModel, EmailStr
from langchain.chains import LLMChain
from langchain_core.prompts import (
    ChatPromptTemplate,
    HumanMessagePromptTemplate,
    MessagesPlaceholder,
)
from langchain_core.messages import SystemMessage
from langchain.chains.conversation.memory import ConversationBufferWindowMemory
from langchain_groq import ChatGroq
from dotenv import load_dotenv
import logging
from fastapi.middleware.cors import CORSMiddleware
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from uuid import uuid4
import aiofiles
import aiofiles.os
import pyrebase
from pydantic_settings import BaseSettings
from passlib.context import CryptContext
import firebase_admin
from firebase_admin import credentials, firestore, auth,initialize_app
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm,HTTPAuthorizationCredentials,HTTPBearer
from fastapi.responses import JSONResponse

# Load environment variables from a .env file
load_dotenv()
cred = credentials.Certificate("eva-firebase.json")

# Firebase Initialization
if not firebase_admin._apps:

    firebase_admin.initialize_app(cred)


#For Firebase JS SDK v7.20.0 and later, measurementId is optional
# firebaseConfig = {
#   "apiKey": "AIzaSyCHvAFw4hvoCGJQq8gk-_98B89kRJAObko",
#   "authDomain": "eva-159c0.firebaseapp.com",
#   "projectId": "eva-159c0",
#   "storageBucket": "eva-159c0.firebasestorage.app",
#   "messagingSenderId": "1030084063155",
#   "appId": "1:1030084063155:web:9969a48495c26911c9e5c6",
#   "measurementId": "G-2W8QQVZK1F",
#  "databaseURL": "",
# }
#
# firebase = pyrebase.initialize_app(firebaseConfig)
# Firestore Initialization
db = firestore.client()

# Password Hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# Configuration using pydantic-settings
class Settings(BaseSettings):
    groq_api_key: str
    document_dir: str = os.path.join(os.path.dirname(__file__), "generated_documents")
    max_file_age_hours: int = 24
    log_level: str = "INFO"

    class Config:
        env_file = ".env"

settings = Settings()

# Initialize logging
logging.basicConfig(level=settings.log_level)

app = FastAPI()
security = HTTPBearer()

# Add CORS middleware
allowed_origins = [
    "http://localhost:5173"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Groq client
model = 'llama-3.2-90b-vision-preview'
groq_chat = ChatGroq(
    groq_api_key=settings.groq_api_key,
    model_name=model,
    temperature=0.0
)

# Setup conversation memory
conversational_memory_length = 20
memory = ConversationBufferWindowMemory(k=conversational_memory_length, memory_key="chat_history", return_messages=True)

system_prompt = (
    'You are Educational Virtual Assistant (EVA) that is an AI-powered platform designed to streamline and enhance '
    'the educational process for educators and students. Act as an interviewer and always respond with a follow-up '
    'question to gather more information from the user, and stop responding questions after giving the final response. Do not provide direct answers, but instead ask follow-up questions '
    'to help the user refine their input. Remember to stay focused on educational topics and assist the user in creating '
    'tailored prompts for their needs. Avoid asking more than 7 questions before generating the final response. Provide the final answer after gathered the minimum details and make sure to provide all the answers as well for the assignments, quizzes or rubrics.'
)

# Initialize conversation chain
conversation = LLMChain(
    llm=groq_chat,
    prompt=ChatPromptTemplate.from_messages(
        [
            SystemMessage(content=system_prompt),
            MessagesPlaceholder(variable_name="chat_history"),
            HumanMessagePromptTemplate.from_template("{human_input}"),
        ]
    ),
    verbose=False,
    memory=memory,
)

# Create a directory for storing Word documents
DOCUMENT_DIR = settings.document_dir
os.makedirs(DOCUMENT_DIR, exist_ok=True)

class Query(BaseModel):
    question: str
    user_id: str

class DocumentRequest(BaseModel):
    user_id: str

class UserSignup(BaseModel):
    name: str
    email: EmailStr
    password: str

class UserLogin(BaseModel):
    email: EmailStr
    password: str

def hash_password(password: str) -> str:
    return pwd_context.hash(password)

def verify_password(plain_password, hashed_password) -> bool:
    return pwd_context.verify(plain_password, hashed_password)

@app.get("/get_user")
async def get_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        # Verify the Firebase ID token
        decoded_token = auth.verify_id_token(credentials.credentials)
        user_id = decoded_token["uid"]

        # Fetch user data from Firestore
        user_ref = db.collection("users").document(user_id)
        user_doc = user_ref.get()

        if not user_doc.exists:
            raise HTTPException(status_code=404, detail="User not found")

        user_data = user_doc.to_dict()
        return {
            "name": user_data["name"],
            "email": user_data["email"],
        }
    except Exception as e:
        raise HTTPException(status_code=401, detail="Invalid token")

@app.post("/signup/")
async def signup(user: UserSignup):
    try:
        # Create user in Firebase Authentication
        user_record = auth.create_user(
            email=user.email,
            password=user.password,
            display_name=user.name
        )

        # Store user data in Firestore
        user_ref = db.collection("users").document(user_record.uid)
        user_ref.set({
            "name": user.name,
            "email": user.email,
            "hashed_password": hash_password(user.password)
        })

        # Generate a Firebase token
        custom_token = auth.create_custom_token(user_record.uid)

        return {
            "message": "User registered successfully",
            "custom_token": custom_token,
            "user_id": user_record.uid
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/login/")
async def login(user: UserLogin):
    try:
        # Sign in the user with Firebase Authentication
        user_record = auth.get_user_by_email(user.email)
        user_ref = db.collection("users").document(user_record.uid)
        user_doc = user_ref.get()

        if not user_doc.exists:
            raise HTTPException(status_code=400, detail="Invalid credentials")

        user_data = user_doc.to_dict()

        if not verify_password(user.password, user_data["hashed_password"]):
            raise HTTPException(status_code=400, detail="Invalid credentials")

        # Generate a Firebase token
        custom_token = auth.create_custom_token(user_record.uid)

        return {
            "message": "Login Successful",
            "custom_token": custom_token,
            "user_id": user_record.uid
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.get("/protected/")
async def protected_route(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        # Verify the Firebase token
        decoded_token = auth.verify_id_token(credentials.credentials)
        user_id = decoded_token["uid"]

        # Fetch user data from Firestore
        user_ref = db.collection("users").document(user_id)
        user_doc = user_ref.get()

        if not user_doc.exists:
            raise HTTPException(status_code=404, detail="User not found")

        user_data = user_doc.to_dict()

        return {
            "message": "Access granted",
            "user": {
                "name": user_data["name"],
                "user_id": user_id
            }
        }
    except Exception as e:
        raise HTTPException(status_code=401, detail="Invalid token")

@app.post("/logout/")
async def logout(request: Request):
    token = request.headers.get("Authorization")
    if not token or not token.startswith("Bearer "):
        raise HTTPException(status_code=400, detail="Invalid token format")

    token = token.split("Bearer ")[1]
    db.collection("token_blacklist").document(token).set({"token": token, "blacklisted_at": datetime.utcnow()})

    return {"message": "Successfully logged out"}

async def create_word_document(user_id: str, questions: str, response: str) -> str:
    os.makedirs(DOCUMENT_DIR, exist_ok=True)
    doc = Document()

    title = doc.add_heading('Educational Virtual Assistant (EVA) Response', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"User ID: {user_id}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading('Questions:', level=1)
    doc.add_paragraph(questions)

    doc.add_heading('Response:', level=1)

    paragraphs = response.split('\n')
    for para in paragraphs:
        if para.strip():
            p = doc.add_paragraph()
            if para.strip().startswith(('•', '-', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                p.style = 'List Bullet'
            p.add_run(para.strip())

    filename = f"eva_response_{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(DOCUMENT_DIR, filename)
    doc.save(filepath)
    return filepath

async def cleanup_old_files(max_age_hours: int = settings.max_file_age_hours):
    current_time = datetime.now()
    for filename in os.listdir(DOCUMENT_DIR):
        file_path = os.path.join(DOCUMENT_DIR, filename)
        file_modified_time = datetime.fromtimestamp(os.path.getmtime(file_path))
        if current_time - file_modified_time > timedelta(hours=max_age_hours):
            await aiofiles.os.remove(file_path)
            logging.info(f"Removed old file: {filename}")

@app.post("/chat/")
async def chat(query: Query):
    user_question = query.question
    user_id = query.user_id

    if not user_question:
        raise HTTPException(status_code=400, detail="Question must not be empty")

    try:
        user_ref = db.collection("users").document(user_id)
        if not user_ref.get().exists:
            raise HTTPException(status_code=404, detail="User not found")

        conversation_ref = user_ref.collection("conversations").document()
        conversation_data = {
            "user_id": user_id,
            "question": user_question,
            "timestamp": datetime.utcnow()
        }
        conversation_ref.set(conversation_data)

        combined_input = user_question
        final_prompt = f"{system_prompt}\n{combined_input}"
        response = conversation.run(human_input=final_prompt)

        conversation_ref.update({"response": response})

        return {"response": response, "status": "success"}

    except Exception as e:
        logging.error(f"Error during chat processing: {e}")
        raise HTTPException(status_code=500, detail="Internal Server Error")

@app.get("/chat_sessions/{user_id}")
async def get_chat_sessions(user_id: str):
    user_ref = db.collection("users").document(user_id)
    if not user_ref.get().exists:
        raise HTTPException(status_code=404, detail="User not found")

    conversations = user_ref.collection("conversations").stream()
    sessions = [{"id": conv.id, **conv.to_dict()} for conv in conversations]

    return {
        "status": "success",
        "sessions": sessions,
    }

@app.get("/chat_session/{session_id}")
async def get_chat_session(session_id: str):
    conversations_ref = db.collection("users").where("conversations", "array_contains", session_id).stream()
    for conv in conversations_ref:
        session = conv.to_dict()
        return {
            "status": "success",
            "session": session,
        }
    raise HTTPException(status_code=404, detail="Session not found")

@app.post("/generate_document/")
async def generate_document(request: DocumentRequest):
    user_id = request.user_id

    user_ref = db.collection("users").document(user_id)
    if not user_ref.get().exists:
        raise HTTPException(status_code=404, detail="User not found")

    last_conversation = user_ref.collection("conversations").order_by("timestamp",
                                                                      direction=firestore.Query.DESCENDING).limit(
        1).stream()
    last_conversation = next(last_conversation, None)
    if not last_conversation:
        raise HTTPException(status_code=400, detail="No conversation available for this user")

    last_conversation_data = last_conversation.to_dict()
    filepath = await create_word_document(user_id, last_conversation_data["question"],
                                          last_conversation_data["response"])

    await cleanup_old_files()

    return FileResponse(filepath, filename=os.path.basename(filepath),
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.post("/new_chat/")
async def new_chat(request: DocumentRequest):
    user_id = request.user_id

    user_ref = db.collection("users").document(user_id)
    if not user_ref.get().exists:
        raise HTTPException(status_code=404, detail="User not found")

    new_conversation_ref = user_ref.collection("conversations").document()
    new_conversation_ref.set({
        "user_id": user_id,
        "timestamp": datetime.utcnow(),
        "messages": []
    })

    return {
        "status": "success",
        "message": "New chat session started",
    }

if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)