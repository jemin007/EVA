import os
from typing import List, Dict, Any
from fastapi import FastAPI, HTTPException, Depends, Request, Security
from fastapi.params import Depends
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel, EmailStr
from langchain.chains import LLMChain
from langchain_core.prompts import (
    ChatPromptTemplate,
    HumanMessagePromptTemplate,
    MessagesPlaceholder,
)
from langchain_core.messages import SystemMessage
from langchain.chains.conversation.memory import ConversationBufferWindowMemory
from langchain_groq import ChatGroq
from dotenv import load_dotenv
import logging
from fastapi.middleware.cors import CORSMiddleware
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from uuid import uuid4
import aiofiles
import aiofiles.os
from passlib.context import CryptContext
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from supabase import create_client, Client
from markdown import markdown
from pydantic import BaseModel
import stripe



class UserIdRequest(BaseModel):
    user_id: str

# Load environment variables from a .env file
load_dotenv()

# Initialize Supabase client
SUPABASE_URL = "https://hfvcvaghaarypglwxktp.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImhmdmN2YWdoYWFyeXBnbHd4a3RwIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NDIyMzEyMzksImV4cCI6MjA1NzgwNzIzOX0.z5SVtyRrmgD-LvSXvOQq4T16DpmRM7sBbptrAfQLvMs"
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# Password Hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# stripe key
stripe.api_key = "sk_test_51RAFGMLGXU8BfuvFJvhqglbARsPIcDsHdkqb4Ft6U8IhFD9rR88ysMYRbfO8bB9QHwyD6PMQvLX5gCoYb1bNqXpg006iz6uAdO"

# Configuration using environment variables
class Settings(BaseModel):
    groq_api_key: str
    document_dir: str = os.path.join(os.path.dirname(__file__), "generated_documents")
    max_file_age_hours: int = 24
    log_level: str = "INFO"


settings = Settings(
    groq_api_key=os.getenv("GROQ_API_KEY"),
    document_dir=os.getenv("DOCUMENT_DIR", os.path.join(os.path.dirname(__file__), "generated_documents")),
    max_file_age_hours=int(os.getenv("MAX_FILE_AGE_HOURS", 24)),
    log_level=os.getenv("LOG_LEVEL", "INFO"),
)

# Initialize logging
logging.basicConfig(level=settings.log_level)

app = FastAPI()
security = HTTPBearer()

# Add CORS middleware
allowed_origins = [
    "http://localhost:5173",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"],

)

# Initialize Groq client
model = 'llama-3.2-90b-vision-preview'
groq_chat = ChatGroq(
    groq_api_key=settings.groq_api_key,
    model_name=model,
    temperature=0.0
)

# Setup conversation memory
conversational_memory_length = 20
memory = ConversationBufferWindowMemory(k=conversational_memory_length, memory_key="chat_history", return_messages=True)

system_prompt = (
    'You are Educational Virtual Assistant (EVA) that is an AI-powered platform designed to streamline and enhance '
    'the educational process for educators and students. Act as an interviewer and always respond with a follow-up '
    'question to gather more information from the user, and stop responding questions after giving the final response. Do not provide direct answers, but instead ask follow-up questions '
    'to help the user refine their input. Remember to stay focused on educational topics and assist the user in creating '
    'tailored prompts for their needs. Avoid asking more than 7 questions before generating the final response. Provide the final answer after gathered the minimum details and make sure to provide all the answers as well for the assignments, quizzes or rubrics.'
)

# Initialize conversation chain
conversation = LLMChain(
    llm=groq_chat,
    prompt=ChatPromptTemplate.from_messages(
        [
            SystemMessage(content=system_prompt),
            MessagesPlaceholder(variable_name="chat_history"),
            HumanMessagePromptTemplate.from_template("{human_input}"),
        ]
    ),
    verbose=False,
    memory=memory,
)

# Create a directory for storing Word documents
DOCUMENT_DIR = settings.document_dir
os.makedirs(DOCUMENT_DIR, exist_ok=True)


# Pydantic Models
class Query(BaseModel):
    question: str
    user_id: str


class DocumentRequest(BaseModel):
    user_id: str


class UserSignup(BaseModel):
    name: str
    email: str
    password: str


class UserLogin(BaseModel):
    email: EmailStr
    password: str


# Utility Functions
def hash_password(password: str) -> str:
    return pwd_context.hash(password)


def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)


class UserResponse:
    def __init__(self, id, name, email):
        self.id = id
        self.name = name
        self.email = email

'''
@app.get("/get_user/")
async def get_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        # Extract the token from the Authorization header
        token = credentials.credentials

        # Verify the token and get the user's info
        response = supabase.auth.get_user(token)

        if not response or "user" not in response:
            raise HTTPException(status_code=401, detail="Invalid token or unauthorized access")

        user_id = response["user"]["id"]

        # Fetch user data from the database
        user_data = supabase.table("users").select("*").eq("id", user_id).execute()
        print(user_data)

        # Check if the user exists
        if not user_data.data or len(user_data.data) == 0:
            raise HTTPException(status_code=404, detail="User not found")

        # Return the user data
        return UserResponse(
            id=user_data.data[0]["id"],
            name=user_data.data[0]["name"],
            email=user_data.data[0]["email"],
        )
    except Exception as e:
        raise HTTPException(status_code=401, detail=str(e))
'''

# Authentication Endpoints
@app.post("/signup/")
async def signup(user: UserSignup):
    try:
        # Create user in Supabase Auth
        auth_response = supabase.auth.sign_up({
            "email": user.email,
            "password": user.password,
        })

        # Store user data in Supabase database
        user_data = {
            "id": auth_response.user.id,
            "name": user.name,
            "email": user.email,
            "hashed_password": hash_password(user.password),
        }
        supabase.table("users").insert(user_data).execute()

        return {
            "message": "User registered successfully",
            "user_id": auth_response.user.id,
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/reset_conversation")
async def reset_conversation(request: UserIdRequest):
    try:
        user_id = request.user_id

        if not user_id:
            raise HTTPException(status_code=400, detail="User ID is required")

        # 1. Completely reset the LangChain memory
        memory.clear()
        memory.chat_memory.clear()  # Clear all message history

        # 2. Delete all related database records
        # Delete all conversation states
        supabase.table("conversation_states").delete().eq("user_id", user_id).execute()

        # Get all sessions for the user
        sessions = supabase.table("chat_sessions") \
            .select("id") \
            .eq("user_id", user_id) \
            .execute()

        # Delete all messages and sessions
        if sessions.data:
            session_ids = [session["id"] for session in sessions.data]
            supabase.table("chat_messages").delete().in_("session_id", session_ids).execute()
            supabase.table("chat_sessions").delete().eq("user_id", user_id).execute()

        # 3. Create fresh session with initial system prompt
        new_session_id = str(uuid4())
        supabase.table("chat_sessions").insert([{
            "id": new_session_id,
            "user_id": user_id,
            "title": "New Chat",
            "created_at": datetime.utcnow().isoformat()
        }]).execute()

        # Initialize with empty conversation state
        supabase.table("conversation_states").insert([{
            "user_id": user_id,
            "conversation_turns": 0,
            "conversation_data": []
        }]).execute()

        # 4. Force a fresh conversation chain
        global conversation
        conversation = LLMChain(
            llm=groq_chat,
            prompt=ChatPromptTemplate.from_messages([
                SystemMessage(content=system_prompt),
                MessagesPlaceholder(variable_name="chat_history"),
                HumanMessagePromptTemplate.from_template("{human_input}"),
            ]),
            verbose=False,
            memory=ConversationBufferWindowMemory(
                k=conversational_memory_length,
                memory_key="chat_history",
                return_messages=True
            )
        )

        return {
            "success": True,
            "new_session_id": new_session_id,
            "message": "Conversation reset successfully"
        }

    except Exception as e:
        logging.error(f"Reset error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# # payment integration
@app.post("/upgrade-plan/")
async def upgrade_plan(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        token = credentials.credentials
        auth_user = supabase.auth.get_user(token)
        user_id = auth_user.user.id

        checkout_session = stripe.checkout.Session.create(
            payment_method_types=['card'],
            line_items=[{'price': 'price_12345', 'quantity': 1}],
            mode='subscription',
            success_url=f"{os.getenv('FRONTEND_URL')}/payment-success",
            cancel_url=f"{os.getenv('FRONTEND_URL')}/pricing",
            metadata={"user_id": user_id}
        )

        return {"url": checkout_session.url}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/user-plan")
async def get_user_plan(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        # Verify token
        user = supabase.auth.get_user(credentials.credentials)
        if not user.user:
            raise HTTPException(status_code=401, detail="Invalid token")

        # Fetch user data
        data = supabase.from_("users") \
            .select("plan_type, responses_generated") \
            .eq("id", user.user.id) \
            .execute()

        if not data.data:
            raise HTTPException(status_code=404, detail="User not found")

        return data.data[0]

    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))
# Payment endpoints
@app.post("/create-checkout-session/")
async def create_checkout_session(
    credentials: HTTPAuthorizationCredentials = Depends(security),
):
    try:
        # 1. Authenticate user
        token = credentials.credentials
        auth_response = supabase.auth.get_user(token)
        user = auth_response.user

        if not user:
            raise HTTPException(status_code=401, detail="Invalid authentication")

        user_id = user.id

        # 2. Get user email from database
        user_query = supabase.from_("users").select("email").eq("id", user_id).execute()
        if not user_query.data or len(user_query.data) == 0:
            raise HTTPException(status_code=404, detail="User not found in database")

        user_email = user_query.data[0]["email"]

        # 3. Create Stripe Checkout Session
        price_id = os.getenv("STRIPE_PRO_PRICE_ID")
        if not price_id:
            raise HTTPException(status_code=500, detail="Stripe price ID is not configured.")

        try:
            checkout_session = stripe.checkout.Session.create(
                payment_method_types=['card'],
                line_items=[{
                    'price': price_id,
                    'quantity': 1,
                }],
                mode='subscription',
                customer_email=user_email,
                metadata={
                    "user_id": user_id,
                    "plan": "pro"
                },
                success_url="http://localhost:5173/payment-success?session_id={CHECKOUT_SESSION_ID}",
                cancel_url="http://localhost:5173/pricing",
            )
        except stripe.error.StripeError as e:
            raise HTTPException(status_code=500, detail=str(e))

        return {"url": checkout_session.url}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    except HTTPException:
        raise  # Re-raise known exceptions
    except stripe.error.StripeError as e:
        logging.error(f"Stripe error: {str(e)}")
        raise HTTPException(status_code=400, detail=str(e.user_message))
    except Exception as e:
        logging.error(f"Unexpected error: {str(e)}")
        raise HTTPException(status_code=500, detail="Internal server error")

@app.post("/create-payment-intent/")
async def create_payment_intent(request: Request):
    data = await request.json()

    intent = stripe.PaymentIntent.create(
        amount=1500,  # $15.00
        currency='usd',
        metadata={
            'plan': data['plan'],
            'user_id': request.user.id
        }
    )

    return {'clientSecret': intent.client_secret}


@app.get("/verify-payment")
async def verify_payment(session_id: str):
    session = stripe.checkout.Session.retrieve(session_id)
    if session.payment_status == "paid":
        # Get user ID from session metadata
        user_id = session.metadata.get("user_id")

        # Return fresh plan data
        user_data = supabase.from_("users") \
            .select("plan_type, responses_generated") \
            .eq("id", user_id) \
            .execute()

        return {
            "status": "success",
            "plan_type": user_data.data[0]["plan_type"],
            "responses_generated": user_data.data[0]["responses_generated"]
        }
    return {"status": "pending"}

# Webhook handler (improved)
@app.post("/stripe-webhook")
async def handle_stripe_webhook(request: Request):
    payload = await request.body()
    sig_header = request.headers.get("stripe-signature")

    try:
        event = stripe.Webhook.construct_event(
            payload,
            sig_header,
            os.getenv("STRIPE_WEBHOOK_SECRET")
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail="Invalid payload")
    except stripe.error.SignatureVerificationError as e:
        raise HTTPException(status_code=400, detail="Invalid signature")

    if event.type == "checkout.session.completed":
        session = event.data.object
        user_id = session.metadata.get("user_id")

        # Update Supabase
        update_response = supabase.table("users") \
            .update({"plan_type": "pro"}) \
            .eq("id", user_id) \
            .execute()

        if not update_response.data:
            raise HTTPException(status_code=400, detail="Failed to update user plan")

    return {"status": "success"}
@app.post("/login/")
async def login(user: UserLogin):
    try:
        # Sign in the user with Supabase
        auth_response = supabase.auth.sign_in_with_password({
            "email": user.email,
            "password": user.password,
        })

        return {
            "message": "Login Successful",
            "access_token": auth_response.session.access_token,
            "user_id": auth_response.user.id,
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/protected/")
async def protected_route(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        # Verify the Supabase token
        token = credentials.credentials
        user = supabase.auth.get_user(token)

        # Fetch user data from Supabase
        user_data = supabase.table("users").select("*").eq("id", user.user.id).execute()
        if not user_data.data:
            raise HTTPException(status_code=404, detail="User not found")

        return {
            "message": "Access granted",
            "user": {
                "name": user_data.data[0]["name"],
                "user_id": user.user.id,
            }
        }
    except Exception as e:
        raise HTTPException(status_code=401, detail="Invalid token")


@app.post("/chat/")
async def chat(query: Query, credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        # Verify authentication
        token = credentials.credentials
        auth_user = supabase.auth.get_user(token)
        user_id = auth_user.user.id

        # Fetch user data with plan info
        user_data = supabase.table("users") \
            .select("plan_type, responses_generated") \
            .eq("id", user_id) \
            .execute()

        if not user_data.data:
            raise HTTPException(status_code=404, detail="User not found")

        user = user_data.data[0]
        question = query.question

        # Check free user limits
        if user["plan_type"] == "free" and user["responses_generated"] >= 2:
            raise HTTPException(
                status_code=403,
                detail="Free plan limit reached. Upgrade to continue chatting."
            )

        # Get conversation state
        state_data = supabase.table("conversation_states") \
            .select("*") \
            .eq("user_id", user_id) \
            .execute()

        is_new_conversation = not state_data.data or state_data.data[0]["conversation_turns"] == 0

        # Get or create session
        session_data = supabase.table("chat_sessions") \
            .select("*") \
            .eq("user_id", user_id) \
            .order("created_at", desc=True) \
            .limit(1) \
            .execute()

        if not session_data.data:
            current_session = {
                "id": str(uuid4()),
                "user_id": user_id,
                "title": question[:50],
                "created_at": datetime.utcnow().isoformat(),
            }
            supabase.table("chat_sessions").insert([current_session]).execute()
        else:
            current_session = session_data.data[0]

        # Save user message
        supabase.table("chat_messages").insert([{
            "session_id": current_session["id"],
            "sender": "user",
            "text": question,
            "timestamp": datetime.utcnow().isoformat(),
        }]).execute()

        # Generate response
        if is_new_conversation:
            response = conversation.run(human_input=question)
        else:
            MAX_HISTORY_LENGTH = 5
            history = state_data.data[0]["conversation_data"][-MAX_HISTORY_LENGTH:]
            combined_input = " ".join([qa["response"] for qa in history])
            response = conversation.run(human_input=combined_input + "\n" + question)

        # Save bot response
        supabase.table("chat_messages").insert([{
            "session_id": current_session["id"],
            "sender": "bot",
            "text": response,
            "timestamp": datetime.utcnow().isoformat(),
        }]).execute()

        # Update conversation state
        if is_new_conversation:
            supabase.table("conversation_states").insert([{
                "user_id": user_id,
                "conversation_turns": 1,
                "conversation_data": [{"response": question}],
            }]).execute()
        else:
            updated_state = {
                "conversation_turns": state_data.data[0]["conversation_turns"] + 1,
                "conversation_data": state_data.data[0]["conversation_data"] + [{"response": question}],
            }
            supabase.table("conversation_states") \
                .update(updated_state) \
                .eq("user_id", user_id) \
                .execute()

        # Increment response count for free users
        if user["plan_type"] == "free":
            supabase.table("users") \
                .update({"responses_generated": user["responses_generated"] + 1}) \
                .eq("id", user_id) \
                .execute()

        html_response = markdown(response)
        return {
            "response": html_response,
            "raw_response": response,
            "remaining_responses": max(0, 2 - (user["responses_generated"] + 1)) if user[
                                                                                        "plan_type"] == "free" else "unlimited",
            "is_new_conversation": is_new_conversation
        }

    except HTTPException as he:
        raise he
    except Exception as e:
        logging.error(f"Chat error: {str(e)}")
        raise HTTPException(status_code=500, detail="Internal server error")

# Document Generation Endpoint
@app.post("/generate_document/")
async def generate_document(request: DocumentRequest):
    user_id = request.user_id

    # Fetch user data from Supabase
    user_data = supabase.table("users").select("*").eq("id", user_id).execute()
    if not user_data.data:
        raise HTTPException(status_code=404, detail="User not found")

    # Fetch the latest conversation
    last_conversation = supabase.table("conversations").select("*").eq("user_id", user_id).order("timestamp",
                                                                                                 desc=True).limit(
        1).execute()
    if not last_conversation.data:
        raise HTTPException(status_code=400, detail="No conversation available for this user")

    last_conversation_data = last_conversation.data[0]
    filepath = await create_word_document(user_id, last_conversation_data["question"],
                                          last_conversation_data["response"])

    await cleanup_old_files()

    return FileResponse(filepath, filename=os.path.basename(filepath),
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# Utility Functions for Document Generation
async def create_word_document(user_id: str, questions: str, response: str) -> str:
    os.makedirs(DOCUMENT_DIR, exist_ok=True)
    doc = Document()

    title = doc.add_heading('Educational Virtual Assistant (EVA) Response', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"User ID: {user_id}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading('Questions:', level=1)
    doc.add_paragraph(questions)

    doc.add_heading('Response:', level=1)

    paragraphs = response.split('\n')
    for para in paragraphs:
        if para.strip():
            p = doc.add_paragraph()
            if para.strip().startswith(('•', '-', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                p.style = 'List Bullet'
            p.add_run(para.strip())

    filename = f"eva_response_{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(DOCUMENT_DIR, filename)
    doc.save(filepath)
    return filepath


async def cleanup_old_files(max_age_hours: int = settings.max_file_age_hours):
    current_time = datetime.now()
    for filename in os.listdir(DOCUMENT_DIR):
        file_path = os.path.join(DOCUMENT_DIR, filename)
        file_modified_time = datetime.fromtimestamp(os.path.getmtime(file_path))
        if current_time - file_modified_time > timedelta(hours=max_age_hours):
            await aiofiles.os.remove(file_path)
            logging.info(f"Removed old file: {filename}")


# Run the Application
if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8080)
