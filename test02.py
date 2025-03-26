import os
from typing import List, Dict, Any
from fastapi import FastAPI, HTTPException, Depends, Request, Security
from fastapi.params import Depends
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel, EmailStr
from langchain.chains import LLMChain
from langchain_core.prompts import (
    ChatPromptTemplate,
    HumanMessagePromptTemplate,
    MessagesPlaceholder,
)
from langchain_core.messages import SystemMessage
from langchain.chains.conversation.memory import ConversationBufferWindowMemory
from langchain_groq import ChatGroq
from dotenv import load_dotenv
import logging
from fastapi.middleware.cors import CORSMiddleware
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from uuid import uuid4
import aiofiles
import aiofiles.os
from passlib.context import CryptContext
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from supabase import create_client, Client
from markdown import markdown

# Load environment variables from a .env file
load_dotenv()

# Initialize Supabase client
SUPABASE_URL = "https://hfvcvaghaarypglwxktp.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImhmdmN2YWdoYWFyeXBnbHd4a3RwIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NDIyMzEyMzksImV4cCI6MjA1NzgwNzIzOX0.z5SVtyRrmgD-LvSXvOQq4T16DpmRM7sBbptrAfQLvMs"
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# Password Hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")


# Configuration using environment variables
class Settings(BaseModel):
    groq_api_key: str
    document_dir: str = os.path.join(os.path.dirname(__file__), "generated_documents")
    max_file_age_hours: int = 24
    log_level: str = "INFO"


settings = Settings(
    groq_api_key=os.getenv("GROQ_API_KEY"),
    document_dir=os.getenv("DOCUMENT_DIR", os.path.join(os.path.dirname(__file__), "generated_documents")),
    max_file_age_hours=int(os.getenv("MAX_FILE_AGE_HOURS", 24)),
    log_level=os.getenv("LOG_LEVEL", "INFO"),
)

# Initialize logging
logging.basicConfig(level=settings.log_level)

app = FastAPI()
security = HTTPBearer()

# Add CORS middleware
allowed_origins = [
    #"http://localhost:5173",
    "http://localhost:5174",

    # "http://localhost:3000",

]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Groq client
model = 'llama-3.2-90b-vision-preview'
groq_chat = ChatGroq(
    groq_api_key=settings.groq_api_key,
    model_name=model,
    temperature=0.1
)

# Setup conversation memory
conversational_memory_length = 20
memory = ConversationBufferWindowMemory(k=conversational_memory_length, memory_key="chat_history", return_messages=True)

system_prompt = """You are EVA, a friendly Academic Content Assistant. Use casual but professional English. Help users create academic materials through conversation, not interrogation.

**Core Principles:**
1. Start with short greeting ("Hi! How can I assist with your academic content?")
2. First ask: "What would you like to create? (Course outline/Lesson plan/Assignment/Rubric/Grading system)" 
3. Ask a maximum of 6 questions to provide final answer
4. For each request:
   - Identify the 2-3 MOST ESSENTIAL requirements needed to start
   - Ask questions NATURALLY within conversation flow ("First, what's the main goal of this lesson?")
   - Suggest common academic standards unless user specifies otherwise
   - Never ask back-to-back questions

**Special Handling:**
- If user provides minimal info: "Would you like me to suggest standard components for [content type]?"
- For vague requests: "Should we focus on [aspect A] or [aspect B] first?"
- Off-topic: "I'm best at academic content - want help structuring something specific?"
- Never say "good morning" repeatedly or use formal greetings after initial message

**Response Style:**
- Use contractions ("Let's", "We'll")
- Keep responses under 2 sentences unless explaining complex concepts
- Add occasional encouragement ("Good choice!", "Clear objective!")
- Always leave conversation open-ended after responses

 Always format responses using Markdown:

**Formatting Rules:**
1. Use headings for sections: `## Heading`
2. Use bullet points for lists:
   - Item 1
   - Item 2
3. Number steps:
   1. First step
   2. Second step
4. Separate paragraphs with blank lines
5. Use `**bold**` for emphasis
6. Code blocks: ```python\nprint("Hello")\n```

Example formatted response:
## Lesson Plan Structure

**Objective:**  
Understand cellular biology

**Key Topics:**
- Cell structure
- Organelle functions
- Cell division

**Activities:**
1. Microscope lab
2. Group discussion
3. Quiz

Never respond with unformatted paragraphs only.
"""

# Initialize conversation chain
conversation = LLMChain(
    llm=groq_chat,
    prompt=ChatPromptTemplate.from_messages(
        [
            SystemMessage(content=system_prompt),
            MessagesPlaceholder(variable_name="chat_history"),
            HumanMessagePromptTemplate.from_template("{human_input}"),
        ]
    ),
    verbose=False,
    memory=memory,
)

# Create a directory for storing Word documents
DOCUMENT_DIR = settings.document_dir
os.makedirs(DOCUMENT_DIR, exist_ok=True)


# Pydantic Models
class Query(BaseModel):
    question: str
    user_id: str


class DocumentRequest(BaseModel):
    user_id: str


class UserSignup(BaseModel):
    name: str
    email: str
    password: str


class UserLogin(BaseModel):
    email: EmailStr
    password: str


# Utility Functions
def hash_password(password: str) -> str:
    return pwd_context.hash(password)


def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)


class UserResponse:
    def __init__(self, id, name, email):
        self.id = id
        self.name = name
        self.email = email

'''
@app.get("/get_user/")
async def get_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        # Extract the token from the Authorization header
        token = credentials.credentials

        # Verify the token and get the user's info
        response = supabase.auth.get_user(token)

        if not response or "user" not in response:
            raise HTTPException(status_code=401, detail="Invalid token or unauthorized access")

        user_id = response["user"]["id"]

        # Fetch user data from the database
        user_data = supabase.table("users").select("*").eq("id", user_id).execute()
        print(user_data)

        # Check if the user exists
        if not user_data.data or len(user_data.data) == 0:
            raise HTTPException(status_code=404, detail="User not found")

        # Return the user data
        return UserResponse(
            id=user_data.data[0]["id"],
            name=user_data.data[0]["name"],
            email=user_data.data[0]["email"],
        )
    except Exception as e:
        raise HTTPException(status_code=401, detail=str(e))
'''

# Authentication Endpoints
@app.post("/signup/")
async def signup(user: UserSignup):
    try:
        # Create user in Supabase Auth
        auth_response = supabase.auth.sign_up({
            "email": user.email,
            "password": user.password,
        })

        # Store user data in Supabase database
        user_data = {
            "id": auth_response.user.id,
            "name": user.name,
            "email": user.email,
            "hashed_password": hash_password(user.password),
        }
        supabase.table("users").insert(user_data).execute()

        return {
            "message": "User registered successfully",
            "user_id": auth_response.user.id,
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/reset_conversation")
async def reset_conversation(query: Query):
    try:
        user_id = query.user_id

        if not user_id:
            raise HTTPException(status_code=400, detail="User ID is required")

        # 1. Get current session ID before deletion
        current_session = supabase.table("chat_sessions") \
            .select("id") \
            .eq("user_id", user_id) \
            .order("created_at", desc=True) \
            .limit(1) \
            .execute()
        
        current_session_id = current_session.data[0]["id"] if current_session.data else None

        # 2. Delete all related data
        supabase.table("conversation_states").delete().eq("user_id", user_id).execute()
        if current_session_id:
            supabase.table("chat_messages").delete().eq("session_id", current_session_id).execute()
            supabase.table("chat_sessions").delete().eq("id", current_session_id).execute()

        # 3. Create fresh session and state
        new_session_id = str(uuid4())
        supabase.table("chat_sessions").insert([{
            "id": new_session_id,
            "user_id": user_id,
            "title": "New Chat",
            "created_at": datetime.utcnow().isoformat()
        }]).execute()

        supabase.table("conversation_states").insert([{
            "user_id": user_id,
            "conversation_turns": 0,
            "conversation_data": []
        }]).execute()

        return {
            "success": True,
            "new_session_id": new_session_id,
            "message": "Conversation reset successfully"
        }

    except Exception as e:
        logging.error(f"Reset error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
    

@app.post("/login/")
async def login(user: UserLogin):
    try:
        # Sign in the user with Supabase
        auth_response = supabase.auth.sign_in_with_password({
            "email": user.email,
            "password": user.password,
        })

        return {
            "message": "Login Successful",
            "access_token": auth_response.session.access_token,
            "user_id": auth_response.user.id,
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/protected/")
async def protected_route(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        # Verify the Supabase token
        token = credentials.credentials
        user = supabase.auth.get_user(token)

        # Fetch user data from Supabase
        user_data = supabase.table("users").select("*").eq("id", user.user.id).execute()
        if not user_data.data:
            raise HTTPException(status_code=404, detail="User not found")

        return {
            "message": "Access granted",
            "user": {
                "name": user_data.data[0]["name"],
                "user_id": user.user.id,
            }
        }
    except Exception as e:
        raise HTTPException(status_code=401, detail="Invalid token")


@app.post("/chat/")
async def chat(query: Query):
    user_question = query.question
    user_id = query.user_id

    if not user_question:
        raise HTTPException(status_code=400, detail="Question must not be empty")

    try:
        # Check if the user exists
        user = supabase.table("users").select("*").eq("id", user_id).execute()
        if not user.data:
            raise HTTPException(status_code=404, detail="User not found")

        # Fetch or initialize conversation state
        state_data = supabase.table("conversation_states").select("*").eq("user_id", user_id).execute()
        if not state_data.data:
            state = {
                "user_id": user_id,
                "conversation_turns": 0,
                "conversation_data": [],
            }
            supabase.table("conversation_states").insert([state]).execute()
        else:
            state = state_data.data[0]

        # Fetch or initialize current chat session
        session_data = supabase.table("chat_sessions").select("*").eq("user_id", user_id).order("created_at",
                                                                                                desc="true").limit(
            1).execute()
        if not session_data.data:
            current_session = {
                "id": str(uuid4()),
                "user_id": user_id,
                "title": user_question[:50],
                "created_at": datetime.utcnow().isoformat(),
            }
            supabase.table("chat_sessions").insert([current_session]).execute()
        else:
            current_session = session_data.data[0]

        # Save the user's message
        supabase.table("chat_messages").insert([{
            "session_id": current_session["id"],
            "sender": "user",
            "text": user_question,
            "timestamp": datetime.utcnow().isoformat(),
        }]).execute()

        # Generate response
        MAX_HISTORY_LENGTH = 5
        combined_input = " ".join([qa["response"] for qa in state["conversation_data"][-MAX_HISTORY_LENGTH:]])
        final_prompt = f"{system_prompt}\n{combined_input}"
        response = conversation.run(human_input=final_prompt)

        # Save the bot's response
        supabase.table("chat_messages").insert([{
            "session_id": current_session["id"],
            "sender": "bot",
            "text": response,
            "timestamp": datetime.utcnow().isoformat(),
        }]).execute()

        # Update conversation state
        updated_state = {
            "conversation_turns": state["conversation_turns"] + 1,
            "conversation_data": state["conversation_data"] + [{"response": user_question}],
        }
        supabase.table("conversation_states").update(updated_state).eq("user_id", user_id).execute()

        html_response = markdown(response)

        return {"response": html_response, "raw_response": response, "status": "success"}

    except HTTPException as e:
        raise e
    except Exception as e:
        logging.error(f"Error during chat processing: {e}")
        raise HTTPException(status_code=500, detail="Internal Server Error")


# Document Generation Endpoint
@app.post("/generate_document/")
async def generate_document(request: DocumentRequest):
    user_id = request.user_id

    # Fetch user data from Supabase
    user_data = supabase.table("users").select("*").eq("id", user_id).execute()
    if not user_data.data:
        raise HTTPException(status_code=404, detail="User not found")

    # Fetch the latest conversation
    last_conversation = supabase.table("conversations").select("*").eq("user_id", user_id).order("timestamp",
                                                                                                 desc=True).limit(
        1).execute()
    if not last_conversation.data:
        raise HTTPException(status_code=400, detail="No conversation available for this user")

    last_conversation_data = last_conversation.data[0]
    filepath = await create_word_document(user_id, last_conversation_data["question"],
                                          last_conversation_data["response"])

    await cleanup_old_files()

    return FileResponse(filepath, filename=os.path.basename(filepath),
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# Utility Functions for Document Generation
async def create_word_document(user_id: str, questions: str, response: str) -> str:
    os.makedirs(DOCUMENT_DIR, exist_ok=True)
    doc = Document()

    title = doc.add_heading('Educational Virtual Assistant (EVA) Response', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"User ID: {user_id}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading('Questions:', level=1)
    doc.add_paragraph(questions)

    doc.add_heading('Response:', level=1)

    paragraphs = response.split('\n')
    for para in paragraphs:
        if para.strip():
            p = doc.add_paragraph()
            if para.strip().startswith(('•', '-', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                p.style = 'List Bullet'
            p.add_run(para.strip())

    filename = f"eva_response_{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(DOCUMENT_DIR, filename)
    doc.save(filepath)
    return filepath


async def cleanup_old_files(max_age_hours: int = settings.max_file_age_hours):
    current_time = datetime.now()
    for filename in os.listdir(DOCUMENT_DIR):
        file_path = os.path.join(DOCUMENT_DIR, filename)
        file_modified_time = datetime.fromtimestamp(os.path.getmtime(file_path))
        if current_time - file_modified_time > timedelta(hours=max_age_hours):
            await aiofiles.os.remove(file_path)
            logging.info(f"Removed old file: {filename}")


# Run the Application
if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)
