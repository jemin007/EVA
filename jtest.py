import os
from typing import List, Dict, Any
from fastapi import FastAPI, HTTPException , Depends
from fastapi.responses import FileResponse
from pydantic import BaseModel, EmailStr
from langchain.chains import LLMChain
from langchain_core.prompts import (
    ChatPromptTemplate,
    HumanMessagePromptTemplate,
    MessagesPlaceholder,
)
from langchain_core.messages import SystemMessage
from langchain.chains.conversation.memory import ConversationBufferWindowMemory
from langchain_groq import ChatGroq
from dotenv import load_dotenv
import logging
from fastapi.middleware.cors import CORSMiddleware
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from uuid import uuid4
import aiofiles
import aiofiles.os
from pydantic_settings import BaseSettings
from passlib.context import CryptContext
from jose import JWTError, jwt
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from sqlalchemy.orm import Session
from fastapi import FastAPI, HTTPException, Depends

from sqlalchemy import create_engine, Column, Integer, String
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
from databases import Database

app = FastAPI()

# Load environment variables from a .env file
load_dotenv()


''' DB Connection '''
DATABASE_URL = "mysql+mysqlconnector://myuser:mypassword@localhost/my_fastapi_db"

Base = declarative_base()

class User(Base):
    __tablename__ = 'users'
    id = Column(Integer, primary_key=True, index=True)
    name = Column(String(50))
    email = Column(String(100), unique=True, index=True)
    hashed_password = Column(String(100))


engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
database = Database(DATABASE_URL)

Base.metadata.create_all(bind=engine)


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()
''' DB Connection '''

# Configuration using pydantic-settings
class Settings(BaseSettings):
    groq_api_key: str
    document_dir: str = os.path.join(os.path.dirname(__file__), "generated_documents")
    max_file_age_hours: int = 24
    log_level: str = "INFO"

    class Config:
        env_file = ".env"


settings = Settings()

# Initialize logging
logging.basicConfig(level=settings.log_level)

# In-memory storage
conversation_states = {}
chat_history = []
chat_sessions = {}
last_responses = {}  # New dictionary to store the last response for each user



# Add CORS middleware
allowed_origins = [
    "http://localhost:3000",
    "http://localhost:5173",
    "https://evatool.ai/"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# Initialize Groq client
#model = 'llama3-70b-8192'
# testing more models
#model = 'llama-3.2-90b-vision-preview'
# model= "llama-3.1-70b-versatile"
model = "llama-3.3-70b-versatile"
#model = 'deepseek-r1-distill-llama-70b
# test with openai
#model="gpt-4o-mini"

groq_chat = ChatGroq(
    groq_api_key=settings.groq_api_key,
    model_name=model
)

# Setup conversation memory
conversational_memory_length = 20
memory = ConversationBufferWindowMemory(k=conversational_memory_length, memory_key="chat_history", return_messages=True)

# Setup system prompt
system_prompt = (
    'You are Educational Virtual Assistant (EVA) that is an AI-powered platform designed to streamline and enhance '
    'the educational process for educators and students. Act as an interviewer and always respond with a follow-up '
    'question to gather more information from the user, and stop responding questions after giving the final response. Do not provide direct answers, but instead ask follow-up questions '
    'to help the user refine their input. Remember to stay focused on educational topics and assist the user in creating '
    'tailored prompts for their needs. Avoid asking more than 7 questions before generating the final response. Provide the final answer after gathered the minimum details and make sure to provide all the answers as well for the assignments, quizzes or rubrics.'
)

# Initialize conversation chain
conversation = LLMChain(
    llm=groq_chat,
    prompt=ChatPromptTemplate.from_messages(
        [
            SystemMessage(content=system_prompt),
            MessagesPlaceholder(variable_name="chat_history"),
            HumanMessagePromptTemplate.from_template("{human_input}"),
        ]
    ),
    verbose=False,
    memory=memory,
)

# Create a directory for storing Word documents
DOCUMENT_DIR = settings.document_dir
os.makedirs(DOCUMENT_DIR, exist_ok=True)

#user signup in memory testing

users_db = {}
# JWT Secret Key & Expiry
SECRET_KEY = "09d25e094faa6ca2556c818166b7a9563b93f7099f6f0f4caa6cf63b88e8d3e7"
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 30
# Password Hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="login")


class Query(BaseModel):
    question: str
    user_id: str


class DocumentRequest(BaseModel):
    user_id: str

# user signup class
class UserSignup(BaseModel):
    name:str
    email:EmailStr
    password:str

# user login class
class UserLogin(BaseModel):
    email: EmailStr
    password: str

# JWT Token Response Model
class Token(BaseModel):
    access_token: str
    token_type: str

# Function to Hash Password
def hash_password(password: str) -> str:
    return pwd_context.hash(password)

#  Function to Verify Password
def verify_password(plain_password, hashed_password) -> bool:
    return pwd_context.verify(plain_password, hashed_password)


# Function to create JWT token
def create_access_token(data: dict, expires_delta: timedelta):
    to_encode = data.copy()
    expire = datetime.utcnow() + expires_delta
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

@app.post("/signup/")
async def signup(user: UserSignup, db: Session = Depends(get_db)):
    db_user = db.query(User).filter(User.email == user.email).first()
    if db_user:
        raise HTTPException(status_code=400, detail="Email already registered")

    hashed_password = hash_password(user.password)
    new_user = User(name=user.name, email=user.email, hashed_password=hashed_password)
    db.add(new_user)
    db.commit()
    db.refresh(new_user)

    access_token = create_access_token(data={"sub": new_user.email},
                                       expires_delta=timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES))

    return {"message": "User registered successfully", "access_token": access_token, "token_type": "bearer"}

# Login Route
@app.post("/login/")
async def login(user: UserLogin, db: Session = Depends(get_db)):
    db_user = db.query(User).filter(User.email == user.email).first()
    if not db_user or not verify_password(user.password, db_user.hashed_password):
        raise HTTPException(status_code=400, detail="Invalid credentials")

    access_token = create_access_token(
        data={"sub": db_user.email},
        expires_delta=timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES),
    )

    return {"message": "Login Successful", "access_token": access_token, "token_type": "bearer"}



# Protected Route
@app.get("/protected/")
async def protected_route(token: str = Depends(oauth2_scheme)):
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        email: str = payload.get("sub")

        if email not in users_db:
            raise HTTPException(status_code=401, detail="Invalid token")

        return {"message": "Access granted", "user": users_db[email]["name"]}
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid token")

async def create_word_document(user_id: str, questions: str, response: str) -> str:
    """Create a Word document with the chat history and response."""
    os.makedirs(DOCUMENT_DIR, exist_ok=True)  # Ensure directory exists
    doc = Document()

    # Add title
    title = doc.add_heading('Educational Virtual Assistant (EVA) Response', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add user information
    doc.add_paragraph(f"User ID: {user_id}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Add questions section
    doc.add_heading('Questions:', level=1)
    doc.add_paragraph(questions)

    # Add response section
    doc.add_heading('Response:', level=1)

    # Split the response into paragraphs and add them to the document
    paragraphs = response.split('\n')
    for para in paragraphs:
        if para.strip():
            p = doc.add_paragraph()
            if para.strip().startswith(('•', '-', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                p.style = 'List Bullet'
            p.add_run(para.strip())

    # Save the document
    filename = f"eva_response_{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(DOCUMENT_DIR, filename)
    doc.save(filepath)
    return filepath


async def cleanup_old_files(max_age_hours: int = settings.max_file_age_hours):
    """Remove files older than the specified age."""
    current_time = datetime.now()
    for filename in os.listdir(DOCUMENT_DIR):
        file_path = os.path.join(DOCUMENT_DIR, filename)
        file_modified_time = datetime.fromtimestamp(os.path.getmtime(file_path))
        if current_time - file_modified_time > timedelta(hours=max_age_hours):
            await aiofiles.os.remove(file_path)
            logging.info(f"Removed old file: {filename}")


@app.post("/chat/")
async def chat(query: Query):
    user_question = query.question
    user_id = query.user_id

    if not user_question:
        raise HTTPException(status_code=400, detail="Question must not be empty")

    try:
        # Retrieve or initialize conversation state
        if user_id not in conversation_states:
            conversation_states[user_id] = {
                "user_id": user_id,
                "conversation_turns": 0,
                "conversation_data": [],
            }

        # Retrieve or initialize chat sessions
        if user_id not in chat_sessions:
            chat_sessions[user_id] = {
                "sessions": [],
                "current_session": {
                    "id": str(uuid4()),  # Ensure UUID is string
                    "title": user_question[:50],
                    "messages": [],
                },
            }

        state = conversation_states[user_id]
        current_session = chat_sessions[user_id]["current_session"]

        # Save the user's response
        state["conversation_data"].append({"response": user_question})
        current_session["messages"].append({"sender": "user", "text": user_question})

        # Generate response
        MAX_HISTORY_LENGTH = 5
        combined_input = " ".join([qa["response"] for qa in state["conversation_data"][-MAX_HISTORY_LENGTH:]])

        #combined_input = " ".join([qa["response"] for qa in state["conversation_data"]])
        final_prompt = f"{system_prompt}\n{combined_input}"
        response = conversation.run(human_input=final_prompt)

        # Save the bot's response
        current_session["messages"].append({"sender": "bot", "text": response})

        # Update conversation state
        state["conversation_turns"] += 1
        conversation_states[user_id] = state

        # Save the last response
        last_responses[user_id] = {
            "questions": combined_input,
            "response": response,
        }

        return {"response": response, "status": "success"}

    except Exception as e:
        logging.error(f"Error during chat processing: {e}")
        raise HTTPException(status_code=500, detail="Internal Server Error")


@app.get("/chat_sessions/{user_id}")
async def get_chat_sessions(user_id: str):
    # Initialize if user doesn't exist
    if user_id not in chat_sessions:
        chat_sessions[user_id] = {
            "sessions": [],
            "current_session": {
                "id": str(uuid4()),
                "title": "New Chat",
                "messages": []
            }
        }

    # Return both historical sessions and current session
    return {
        "status": "success",
        "sessions": [
            *chat_sessions[user_id]["sessions"],
            chat_sessions[user_id]["current_session"]
        ],
    }


@app.get("/chat_session/{session_id}")
async def get_chat_session(session_id: str):  # Changed from float to str
    for user_id, sessions in chat_sessions.items():
        # Check both current session and archived sessions
        all_sessions = sessions["sessions"] + [sessions["current_session"]]
        for session in all_sessions:
            if str(session["id"]) == session_id:
                return {
                    "status": "success",
                    "session": session,
                }
    raise HTTPException(status_code=404, detail="Session not found")


@app.post("/generate_document/")
async def generate_document(request: DocumentRequest):
    user_id = request.user_id

    if user_id not in last_responses:
        raise HTTPException(status_code=400, detail="No response available for this user")

    try:
        last_response = last_responses[user_id]
        filepath = await create_word_document(user_id, last_response["questions"], last_response["response"])

        # Cleanup old files
        await cleanup_old_files()

        return FileResponse(filepath, filename=os.path.basename(filepath),
                            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    except FileNotFoundError as e:
        logging.error(f"File not found: {e}")
        raise HTTPException(status_code=404, detail="File not found")
    except Exception as e:
        logging.error(f"Error generating document: {e}")
        raise HTTPException(status_code=500, detail="Error generating document")


@app.post("/new_chat/")
async def new_chat(request: DocumentRequest):
    user_id = request.user_id

    try:
        # Initialize chat_sessions for new users
        if user_id not in chat_sessions:
            chat_sessions[user_id] = {
                "sessions": [],
                "current_session": {
                    "id": str(uuid4()),  # Generate a unique session ID
                    "title": "New Chat",
                    "messages": [],
                },
            }
        else:
            # Archive current session if it has messages
            current_session = chat_sessions[user_id]["current_session"]
            if current_session["messages"]:
                chat_sessions[user_id]["sessions"].append(current_session)

            # Create new session
            chat_sessions[user_id]["current_session"] = {
                "id": str(uuid4()),
                "title": "New Chat",
                "messages": [],
            }

        # Reset the conversation state for the user
        conversation_states[user_id] = {
            "user_id": user_id,
            "conversation_turns": 0,
            "conversation_data": [],
        }

        # Log the new chat session
        logging.info(f"New chat session started for user {user_id}")
        #last_responses.pop(user_id, None)  # Clear previous response data


        return {
            "status": "success",
            "message": "New chat session started",
        }

    except Exception as e:
        logging.error(f"Error starting new chat: {e}")
        raise HTTPException(status_code=500, detail="Error starting new chat")


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)